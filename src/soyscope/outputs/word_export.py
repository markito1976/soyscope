"""Generate Word reports from the SoyScope database using python-docx."""

from __future__ import annotations

import json
import logging
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from ..db import Database

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Styling helpers
# ---------------------------------------------------------------------------

_HEADER_BG = RGBColor(0x4F, 0x81, 0xBD)  # steel-blue header row
_HEADER_FG = RGBColor(0xFF, 0xFF, 0xFF)
_ALT_ROW_BG = RGBColor(0xD9, 0xE2, 0xF3)  # light blue alternate row
_WHITE_BG = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_shading(cell, color: RGBColor) -> None:
    """Apply background shading to a table cell."""
    from docx.oxml.ns import qn
    from lxml import etree

    shading = etree.SubElement(
        cell._element.get_or_add_tcPr(), qn("w:shd")
    )
    shading.set(qn("w:fill"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    shading.set(qn("w:val"), "clear")


def _style_header_row(row, col_count: int) -> None:
    """Apply header styling (background + white bold text) to a table row."""
    for idx in range(col_count):
        cell = row.cells[idx]
        _set_cell_shading(cell, _HEADER_BG)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = _HEADER_FG
                run.font.bold = True
                run.font.size = Pt(10)


def _apply_alternating_rows(table, col_count: int) -> None:
    """Shade data rows with alternating colours (skip header at index 0)."""
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            continue  # header already styled
        bg = _ALT_ROW_BG if row_idx % 2 == 0 else _WHITE_BG
        for c in range(col_count):
            _set_cell_shading(row.cells[c], bg)


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    """Insert a formatted table with header row and alternating colours."""
    col_count = len(headers)
    table = doc.add_table(rows=1, cols=col_count)

    # Try to apply a built-in style; fall back gracefully.
    try:
        table.style = "Light List Accent 1"
    except KeyError:
        pass

    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header row
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
    _style_header_row(table.rows[0], col_count)

    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row_data):
            row_cells[i].text = str(value) if value is not None else ""

    _apply_alternating_rows(table, col_count)


def _truncate(text: str | None, length: int = 120) -> str:
    """Return *text* truncated to *length* characters with an ellipsis."""
    if not text:
        return ""
    if len(text) <= length:
        return text
    return text[: length - 3] + "..."


# ---------------------------------------------------------------------------
# Main exporter
# ---------------------------------------------------------------------------


class WordExporter:
    """Generate a formatted Word (.docx) report from SoyScope data."""

    def __init__(self, db: Database, output_dir: Path) -> None:
        self.db = db
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    # ------------------------------------------------------------------ #
    # Public API
    # ------------------------------------------------------------------ #

    def export(self, filename: str | None = None) -> Path:
        """Build the full Word report and save it to *output_dir*.

        Parameters
        ----------
        filename:
            Optional filename override.  Defaults to
            ``soyscope_report_{YYYY-MM-DD}.docx``.

        Returns
        -------
        Path
            The absolute path to the generated ``.docx`` file.
        """
        today = datetime.now()
        date_str = today.strftime("%Y-%m-%d")

        if filename is None:
            filename = f"soyscope_report_{date_str}.docx"

        doc = Document()

        # Gather data once so every section reuses it.
        stats = self.db.get_stats()

        self._add_title_page(doc, today)
        self._add_executive_summary(doc, stats)
        self._add_database_overview(doc, stats)
        self._add_sector_analysis(doc, stats)
        self._add_top_novel_applications(doc)
        self._add_commercial_applications(doc)
        self._add_timeline_trends(doc, stats)

        out_path = self.output_dir / filename
        doc.save(str(out_path))
        logger.info("Word report saved to %s", out_path)
        return out_path

    # ------------------------------------------------------------------ #
    # Private section builders
    # ------------------------------------------------------------------ #

    def _add_title_page(self, doc: Document, today: datetime) -> None:
        """Title page with centred title, subtitle and date."""
        # Add some blank paragraphs to push the title down
        for _ in range(6):
            doc.add_paragraph("")

        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("SoyScope: Industrial Soy Uses Report")
        run.font.size = Pt(28)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = subtitle_para.add_run(
            "Comprehensive Database of Industrial Soybean Applications"
        )
        sub_run.font.size = Pt(16)
        sub_run.font.color.rgb = RGBColor(0x4F, 0x81, 0xBD)

        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(today.strftime("%B %d, %Y"))
        date_run.font.size = Pt(14)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_executive_summary(self, doc: Document, stats: dict) -> None:
        """Section 1 -- Executive Summary."""
        doc.add_heading("Executive Summary", level=1)

        total = stats.get("total_findings", 0)
        sectors_count = stats.get("total_sectors", 0)
        derivatives_count = stats.get("total_derivatives", 0)

        # Date range
        by_year = stats.get("by_year", {})
        if by_year:
            sorted_years = sorted(by_year.keys())
            date_range = f"{sorted_years[0]} - {sorted_years[-1]}"
        else:
            date_range = "N/A"

        doc.add_paragraph(
            f"This report summarises the SoyScope database which currently "
            f"contains {total:,} findings spanning {date_range}."
        )
        doc.add_paragraph(
            f"The database tracks {sectors_count} industry sectors and "
            f"{derivatives_count} soy derivatives."
        )

        # Top 3 sectors by finding count
        top_sectors = self._query_top_sectors(limit=3)
        if top_sectors:
            doc.add_heading("Top 3 Sectors by Finding Count", level=2)
            for name, count in top_sectors:
                doc.add_paragraph(
                    f"{name}: {count:,} findings", style="List Bullet"
                )

        # Top 3 most novel applications
        top_novel = self._query_top_novel(limit=3)
        if top_novel:
            doc.add_heading("Top 3 Most Novel Applications", level=2)
            for row in top_novel:
                title = row.get("title", "Untitled")
                score = row.get("novelty_score")
                score_str = f" (novelty {score:.2f})" if score is not None else ""
                doc.add_paragraph(
                    f"{title}{score_str}", style="List Bullet"
                )

        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_database_overview(self, doc: Document, stats: dict) -> None:
        """Section 2 -- Database Overview table."""
        doc.add_heading("Database Overview", level=1)

        total = stats.get("total_findings", 0)

        # By source API breakdown
        by_source = stats.get("by_source", {})
        source_breakdown = ", ".join(
            f"{src}: {cnt:,}" for src, cnt in by_source.items()
        ) or "N/A"

        # By source type breakdown
        by_type = stats.get("by_type", {})
        type_breakdown = ", ".join(
            f"{t}: {cnt:,}" for t, cnt in by_type.items()
        ) or "N/A"

        # Enrichment coverage
        enr_catalog = stats.get("enrichment_catalog", 0)
        enr_summary = stats.get("enrichment_summary", 0)
        enr_deep = stats.get("enrichment_deep", 0)
        enrichment_text = (
            f"Catalog: {enr_catalog:,}, Summary: {enr_summary:,}, "
            f"Deep: {enr_deep:,}"
        )

        headers = ["Metric", "Value"]
        rows = [
            ["Total Findings", f"{total:,}"],
            ["By Source API", source_breakdown],
            ["By Source Type", type_breakdown],
            ["Enrichment Coverage", enrichment_text],
        ]

        _add_table(doc, headers, rows)
        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_sector_analysis(self, doc: Document, stats: dict) -> None:
        """Section 3 -- Sector Analysis (top 10 sectors)."""
        doc.add_heading("Sector Analysis", level=1)

        top_sectors = self._query_top_sectors(limit=10)
        if not top_sectors:
            doc.add_paragraph("No sector data available.")
            doc.add_page_break()
            return

        for sector_name, count in top_sectors:
            doc.add_heading(sector_name, level=2)
            doc.add_paragraph(f"Findings in this sector: {count:,}")

            # Top 5 findings in the sector
            findings = self._query_findings_by_sector(sector_name, limit=5)
            for f in findings:
                title = f.get("title", "Untitled")
                year = f.get("year", "N/A")
                summary = _truncate(
                    f.get("ai_summary") or f.get("abstract"), 200
                )
                bullet_text = f"{title} ({year})"
                if summary:
                    bullet_text += f" -- {summary}"
                doc.add_paragraph(bullet_text, style="List Bullet")

        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_top_novel_applications(self, doc: Document) -> None:
        """Section 4 -- Top 20 highest-novelty findings."""
        doc.add_heading("Top Novel Applications", level=1)

        rows_data = self._query_top_novel(limit=20)
        if not rows_data:
            doc.add_paragraph("No enriched findings with novelty scores available.")
            doc.add_page_break()
            return

        headers = ["Rank", "Title", "Year", "Novelty Score", "TRL", "Summary"]
        rows = []
        for rank, row in enumerate(rows_data, start=1):
            title = _truncate(row.get("title", ""), 60)
            year = str(row.get("year", "N/A"))
            novelty = (
                f"{row['novelty_score']:.2f}"
                if row.get("novelty_score") is not None
                else "N/A"
            )
            trl = str(row.get("trl_estimate", "N/A"))
            summary = _truncate(
                row.get("ai_summary") or row.get("abstract"), 80
            )
            rows.append([str(rank), title, year, novelty, trl, summary])

        _add_table(doc, headers, rows)
        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_commercial_applications(self, doc: Document) -> None:
        """Section 5 -- Findings with commercial / scaling / mature status."""
        doc.add_heading("Commercial Applications", level=1)

        commercial_rows = self._query_commercial_findings()
        if not commercial_rows:
            doc.add_paragraph(
                "No findings with commercial, scaling, or mature status."
            )
            doc.add_page_break()
            return

        headers = ["Title", "Year", "Status", "TRL", "Summary"]
        rows = []
        for row in commercial_rows:
            title = _truncate(row.get("title", ""), 60)
            year = str(row.get("year", "N/A"))
            status = row.get("commercialization_status", "N/A")
            trl = str(row.get("trl_estimate", "N/A"))
            summary = _truncate(
                row.get("ai_summary") or row.get("abstract"), 100
            )
            rows.append([title, year, status, trl, summary])

        _add_table(doc, headers, rows)
        doc.add_page_break()

    # ------------------------------------------------------------------ #

    def _add_timeline_trends(self, doc: Document, stats: dict) -> None:
        """Section 6 -- Findings per year."""
        doc.add_heading("Timeline Trends", level=1)

        by_year = stats.get("by_year", {})
        if not by_year:
            doc.add_paragraph("No year data available.")
            return

        headers = ["Year", "Findings"]
        rows = [
            [str(year), f"{count:,}"]
            for year, count in sorted(by_year.items())
        ]

        _add_table(doc, headers, rows)

    # ------------------------------------------------------------------ #
    # Database query helpers
    # ------------------------------------------------------------------ #

    def _query_top_sectors(self, limit: int = 10) -> list[tuple[str, int]]:
        """Return ``[(sector_name, finding_count), ...]`` sorted descending."""
        with self.db.connect() as conn:
            rows = conn.execute(
                """
                SELECT s.name, COUNT(fs.finding_id) AS cnt
                FROM sectors s
                JOIN finding_sectors fs ON s.id = fs.sector_id
                GROUP BY s.name
                ORDER BY cnt DESC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()
            return [(r[0], r[1]) for r in rows]

    def _query_top_novel(self, limit: int = 20) -> list[dict]:
        """Return findings joined with enrichments ordered by novelty score."""
        with self.db.connect() as conn:
            rows = conn.execute(
                """
                SELECT f.title, f.year, f.abstract,
                       e.novelty_score, e.trl_estimate,
                       e.ai_summary, e.commercialization_status
                FROM enrichments e
                JOIN findings f ON e.finding_id = f.id
                WHERE e.novelty_score IS NOT NULL
                ORDER BY e.novelty_score DESC
                LIMIT ?
                """,
                (limit,),
            ).fetchall()
            return [dict(r) for r in rows]

    def _query_findings_by_sector(
        self, sector_name: str, limit: int = 5
    ) -> list[dict]:
        """Return the top *limit* findings for a given sector name."""
        with self.db.connect() as conn:
            rows = conn.execute(
                """
                SELECT f.title, f.year, f.abstract, e.ai_summary
                FROM findings f
                JOIN finding_sectors fs ON f.id = fs.finding_id
                JOIN sectors s ON fs.sector_id = s.id
                LEFT JOIN enrichments e ON f.id = e.finding_id
                WHERE s.name = ?
                ORDER BY f.year DESC, f.id DESC
                LIMIT ?
                """,
                (sector_name, limit),
            ).fetchall()
            return [dict(r) for r in rows]

    def _query_commercial_findings(self) -> list[dict]:
        """Return findings whose commercialization status is commercial, scaling, or mature."""
        with self.db.connect() as conn:
            rows = conn.execute(
                """
                SELECT f.title, f.year, f.abstract,
                       e.commercialization_status, e.trl_estimate, e.ai_summary
                FROM enrichments e
                JOIN findings f ON e.finding_id = f.id
                WHERE e.commercialization_status IN ('commercial', 'scaling', 'mature')
                ORDER BY e.trl_estimate DESC, f.year DESC
                """
            ).fetchall()
            return [dict(r) for r in rows]
